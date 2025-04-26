#python-docx をinstallすること
#deep_translator をinstallすること







import openpyxl
from deep_translator import GoogleTranslator

def tranlater(excel_file='input.xlsx', output_file_name='input_transrated.xlsx'):

    # エクセルファイルを読み込む
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active
        
    #英語質問から日本語質問へ翻訳
    # 翻訳器を初期化
    translator = GoogleTranslator(source='en', target='ja')
        
    # 2行目以降のセルを処理
    for row in ws.iter_rows(min_row=2):
        cell_a = row[0].value or ""   #英語質問 
        cell_b = row[1]               #日本語質問
        if cell_b.value == None: #空白の時だけ翻訳後の文章を書き込む
            translated_text = translator.translate(cell_a)
            cell_b.value = translated_text
    

    #英語回答から日本語回答へ翻訳
    # 翻訳器を初期化
    translator = GoogleTranslator(source='ja', target='en')
        
    # 2行目以降のセルを処理
    for row in ws.iter_rows(min_row=2):
        cell_c = row[2].value or ""   #日本語回答
        cell_d = row[3]               #英語回答
        if cell_d.value == None: #空白の時だけ翻訳後の文章を書き込む
            translated_text = translator.translate(cell_c)
            cell_d.value = translated_text
    
    # エクセルファイルを保存    
    wb.save(output_file_name)
    print("翻訳内容を保存しました。")


tranlater(excel_file='input.xlsx', output_file_name='input_transrated.xlsx')














import openpyxl
from docx import Document
from docx.shared import RGBColor, Inches
from difflib import ndiff
import os
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def make_responce_docx(excel_file = 'input_transrated.xlsx', output_file_name = 'output.docx'):
    
    # エクセルファイルを読み込む
    wb = openpyxl.load_workbook(excel_file)
    ws = wb.active

    # 新しいWordドキュメントを作成
    doc = Document()



    #タイトルの作成
    p = doc.add_paragraph("Responses to reviewers")

    # 2行目以降のセルを処理
    for row in ws.iter_rows(min_row=2):
    
        #セルが空白の場合に空文字列を使用することで、エラーを回避
        cell_a = row[0].value or "" #英語質問 
        cell_b = row[1].value or "" #日本語質問
        cell_c = row[2].value or "" #日本語回答
        cell_d = row[3].value or "" #英語回答
        cell_e = row[4].value or "" #元文章
        cell_f = row[5].value or "" #修正後文章   
        cell_g = row[6].value or "" #ページ等
        cell_h = row[7].value or "" #修正前画像パス
        cell_i = row[8].value or "" #修正後画像パス

        
        
        # 各行ごとに改ページする
        if not cell_a == "":
            doc.add_page_break()  


        #質問
        # A列の内容を見出し2で書き込む（改行あり）
        if not cell_a == "":
            for part in cell_a.split('\n'):
                doc.add_heading(part, level=2)
            p = doc.add_paragraph()

        # A列の内容を太文字で書き込む
        if not cell_a == "":
            p = doc.add_paragraph()
            #各部分を個別にadd_runで追加し、run.add_break()で改行を追加することで、元のテキストの改行を再現
            for part in cell_a.split('\n'):
                run = p.add_run(part)
                run.bold = True #太文字
                run.add_break()

        #!!!下書き限定!!!!!
        # 改行して、B列の内容を太文字で書き込む
        if not cell_b == "":
            p = doc.add_paragraph()
            for part in cell_b.split('\n'):
                run = p.add_run(part)
                run.bold = True
                run.add_break()




        #回答
        ##!!!下書き限定!!!!!
        ## 1行空けて、C列の内容を普通の文字で書き込む
        #if not cell_c == "":
        #    p = doc.add_paragraph()
        #    for part in cell_c.split('\n'):
        #        run = p.add_run(part)
        #        run.add_break()

        ## 1行空けて、D列の内容を普通の文字で書き込む（改行ver）
        #if not cell_d == "":
        #    p = doc.add_paragraph()
        #    for part in cell_d.split('\n'):
        #        run = p.add_run(part)
        #        run.add_break()

        # 1行空けて、D列の内容を普通の文字で書き込む(段落改行ver)
        if not cell_d == "":
            for part in cell_d.split('\n'):
                p = doc.add_paragraph(part)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #両端揃えに設定
                run = p.add_run()
            p = doc.add_paragraph()




        #元原稿
        # 一行空けて、[Original manuscript]という文字を書く
        p = doc.add_paragraph("[Original manuscript]")

        ## 改行して、E列の内容を書き込む（改行ver）
        #p = doc.add_paragraph()
        #for part in cell_e.split('\n'):
        #    run = p.add_run(part)
        #    run.add_break()

        # 改行して、E列の内容を書き込む(段落改行ver)
        for part in cell_e.split('\n'):
            p = doc.add_paragraph(part)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #両端揃えに設定
            run = p.add_run()
        p = doc.add_paragraph()





        #修正後原稿
        # 一行空けて、[Revised manuscript]という文字を書き、続けてG列の内容（ページ数）を書く
        p = doc.add_paragraph("[Revised manuscript]")
        run = p.add_run(cell_g)

        ## 改行して差分を計算して書き込む（改行なしver）
        #diff = list(ndiff(cell_d.split(), cell_e.split()))
        
        # 改行して差分を計算して書き込む（改行ありver）    
        cell_e_temp = cell_e.replace('\n', ' <br> ') # 改行文字を一時的に別の文字列に置き換える
        cell_f_temp = cell_f.replace('\n', ' <br> ') # 改行文字を一時的に別の文字列に置き換える 
        diff = list(ndiff(cell_e_temp.split(), cell_f_temp.split()))

        p = doc.add_paragraph()   
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #両端揃えに設定
        for word in diff:        
            #置き換えた改行文字が来たら改行する    
            if word.endswith('+ <br>') or word.startswith('  <br>'): #改行が追加      
                p = doc.add_paragraph() #改段落
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY #両端揃えに設定
            elif word.startswith('- <br>'): #改行が削除
                continue
            elif word.startswith('- '): #消えた文章
                run = p.add_run(word[2:] + ' ')
                run.font.color.rgb = RGBColor(255, 0, 0) #赤色
                run.font.strike = True #取り消し線
            elif word.startswith('+ '): #追加された文章
                run = p.add_run(word[2:] + ' ')
                run.font.color.rgb = RGBColor(0, 0, 255) #青色
            elif word.startswith('  '): #変化のない文章
                run = p.add_run(word[2:] + ' ')
                run.font.color.rgb = RGBColor(0, 0, 0) #黒色            
        p = doc.add_paragraph() #改段落
        p = doc.add_paragraph() #改段落




        # 画像の挿入
        if (not cell_h == "") and os.path.isfile(cell_h):
            p = doc.add_paragraph("[Original]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #中央ぞろえ
            p.add_run().add_break()  # 改行を追加
            p.add_run().add_picture(cell_h, width=Inches(5)) #画像サイズを指定
            p.add_run().add_break()  # 改行を追加


        if (not cell_i == "") and os.path.isfile(cell_i):        
            p = doc.add_paragraph("[Revised]")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run().add_break()  # 改行を追加
            p.add_run().add_picture(cell_i, width=Inches(5)) #画像サイズを指定



    # docxファイルへの出力
    doc.save(output_file_name)
    print("回答書が", output_file_name, "に保存されました。")





make_responce_docx(excel_file = 'input_transrated.xlsx', output_file_name = 'output.docx')








